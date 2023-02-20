'''
目标：
    用一个txt文件存入邀请人列表, 每个邀请人单独占一行
    自己制作一个邀请函word样式
    用python-docx打开该文档
    每份邀请函添加一个邀请人
    每份邀请函单独占一页
    只需打开一份word,就能打印所有邀请函
技巧：
    读取姓名列表， 名字结构比较重要: 一行一个名字.并去掉换行符
        --fileObj = open(txt_path,'r')
        --txt_content = fileObj.readlines()
        --for invit_name in txt_content:
            --write_guest = invit_name[ :-1] # 去掉换行符
            --guests_list.append(write_guest)
    存入一个列表中， 列表的长度就是制作邀请函的个数
    学会如何在WORD中自定义格式(style): 
        打开一个空白WORD, 
        然后键入 ctrl+alt+shift+s, 
        保存这个sytle, 记住： 仅对该WORD文件有效!!
    在生成的段落中调用该自定义格式(style): 
        --paraObj1 = doc.add_paragraph('It would be a pleasure to have you')
        --paraObj1.style = 'Style1'
    在最后一行的生成内容后， 加入换页符: 
        --paraObj6 = doc.add_paragraph(' ').runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
'''

import docx
from pathlib import Path

txt_path = Path(r'D:\03_program\python\guests_card.txt')
docx_path = Path(r'D:\03_program\python\guests_card.docx')

guests_list = []

fileObj = open(txt_path,'r')
txt_content = fileObj.readlines()
for invit_name in txt_content:
    write_guest = invit_name[ :-1] # 去掉换行符
    guests_list.append(write_guest)
    # print(guests_list)

fileObj.close()

doc = docx.Document(docx_path)
for guestNum in range(len(guests_list)):
    paraObj1 = doc.add_paragraph('It would be a pleasure to have you')
    paraObj1.style = 'Style1'
    paraObj2 = doc.add_paragraph(guests_list[guestNum])
    paraObj2.style = 'Heading 1'
    paraObj3 = doc.add_paragraph('at address…(TBD) to attend the Wedding Celebrate of')
    paraObj3.style = 'Style1'
    paraObj4 = doc.add_paragraph('Feb.28th')
    paraObj4.style = 'Heading 2'
    paraObj5 = doc.add_paragraph('at 7 O\'clock')
    paraObj5.style = 'Style1'
    paraObj6 = doc.add_paragraph(' ').runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

doc.save('guests_invitation.docx')